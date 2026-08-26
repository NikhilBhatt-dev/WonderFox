# -*- coding: utf-8 -*-
"""Builds WonderFox_Interview_Preparation.docx from the real project codebase."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0xE7, 0x5D, 0x2E)      # WonderFox orange
DARK = RGBColor(0x1F, 0x29, 0x37)
BODY = RGBColor(0x33, 0x33, 0x33)
CODE_BG = "F2F2F2"

doc = Document()

# ---------- base + heading styles ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

for name, size, color in [
    ("Heading 1", 18, ACCENT),
    ("Heading 2", 14, ACCENT),
    ("Heading 3", 12, DARK),
]:
    st = doc.styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color

# ---------- page setup ----------
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# ---------- footer page numbers ----------
def add_footer_pagenum(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve"); instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
    p.add_run("  |  WonderFox — Interview Preparation Guide")
    for r in p.runs:
        r.font.size = Pt(9); r.font.color.rgb = DARK

add_footer_pagenum(doc.sections[0])

# ---------- helper functions ----------
def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold
    return p

def rich(parts):
    """parts: list of (text, bold) tuples in one paragraph."""
    p = doc.add_paragraph()
    for text, b in parts:
        r = p.add_run(text); r.bold = b
    return p

def bullets(items, level=1):
    for it in items:
        p = doc.add_paragraph(style="List Bullet" if level == 1 else "List Bullet 2")
        p.add_run(it)

def numbered(items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(it)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(htxt); r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    # ensure header repeat
    trPr = t.rows[0]._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader"); tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)
    return t

def code(text):
    lines = text.split("\n")
    for ln in lines:
        p = doc.add_paragraph()
        r = p.add_run(ln)
        r.font.name = "Consolas"; r.font.size = Pt(9)
        # monospace via rPr
        rpr = r._element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts"); rFonts.set(qn("w:ascii"), "Consolas"); rFonts.set(qn("w:hAnsi"), "Consolas")
        rpr.append(rFonts)
        # shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), CODE_BG)
        pPr.append(shd)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

def spacer():
    doc.add_paragraph()

# =====================================================================
# TITLE PAGE
# =====================================================================
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n\n"); r.font.size = Pt(34)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\U0001F43A  WonderFox"); r.font.size = Pt(40); r.bold = True; r.font.color.rgb = ACCENT
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Soft Toy E-Commerce Store"); r.font.size = Pt(18); r.font.color.rgb = DARK
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Complete Interview Preparation Guide"); r.font.size = Pt(24); r.bold = True; r.font.color.rgb = ACCENT
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Prepared for CSE / BTech Web-Development Technical Interviews"); r.font.size = Pt(12); r.italic = True
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Stack: React + Vite + Tailwind + Node.js + Express + MongoDB"); r.font.size = Pt(12); r.font.color.rgb = BODY
doc.add_page_break()

# =====================================================================
# TABLE OF CONTENTS
# =====================================================================
h1("Table of Contents")
toc_items = [
    ("1.  Project Overview", "1"),
    ("2.  Complete Tech Stack", "2"),
    ("3.  Project Architecture", "3"),
    ("4.  Folder Structure", "4"),
    ("5.  Complete User Flow", "5"),
    ("6.  Important Features - Interview Deep Dive", "6"),
    ("7.  Database Design", "7"),
    ("8.  API Documentation", "8"),
    ("9.  Authentication & Authorization", "9"),
    ("10. E-Commerce Logic", "10"),
    ("11. Important React Concepts Used", "11"),
    ("12. Important JavaScript Concepts Used", "12"),
    ("13. Backend Concepts", "13"),
    ("14. MongoDB & Mongoose Interview Questions", "14"),
    ("15. Most Likely Interview Questions", "15"),
    ("16. \"Tell Me About Your Project\"", "16"),
    ("17. \"What Was Your Contribution?\"", "17"),
    ("18. Challenges Faced", "18"),
    ("19. Technical Decisions", "19"),
    ("20. Performance & Optimization", "20"),
    ("21. Security", "21"),
    ("22. Testing & Debugging", "22"),
    ("23. Deployment", "23"),
    ("24. Scenario-Based Questions", "24"),
    ("25. Code-Level Questions", "25"),
    ("26. Project Weaknesses / Improvements", "26"),
    ("27. Rapid Revision Sheet", "27"),
    ("28. Final Interview Script", "28"),
]
for title, num in toc_items:
    p = doc.add_paragraph()
    r = p.add_run(title); r.font.color.rgb = DARK
    p.add_run("\t\t........ page " + num)
doc.add_page_break()

# =====================================================================
# SECTION 1 - PROJECT OVERVIEW
# =====================================================================
h1("1. Project Overview")

h2("1.1 What is WonderFox?")
para("WonderFox is a full-stack e-commerce web application that sells cute and premium soft toys (teddy bears, plush toys) for kids. It is a \"toy shop\" with three connected apps living in one repo: a customer **storefront** (frontend), a **REST API** (backend), and an **admin dashboard** (admin).")
bullets([
    "frontend/ - the customer storefront (React + Vite + Tailwind). This is what shoppers see.",
    "backend/ - a Node.js + Express + MongoDB REST API that serves both UIs.",
    "admin/ - a separate React panel the owner uses to manage products, categories, orders, users, blogs and newsletters.",
])

h2("1.2 Purpose of the Project")
para("WonderFox demonstrates a complete, production-shaped MERN e-commerce system. It is not a throwaway tutorial clone. It has real business features: a product catalogue, categories, cart, wishlist, guest-cart persistence, order placement (Cash on Delivery + Razorpay), payment signature verification, an admin dashboard, blogs, newsletter email, and a contact form that emails via SMTP.")

h2("1.3 Target Users")
bullets([
    "Customers: people buying soft toys for kids or as gifts.",
    "Admin / store owner: manages the catalogue, orders, users, blogs.",
    "Guest shoppers: can add to cart and wishlist without an account (localStorage), then log in to convert.",
])

h2("1.4 Main Business Problem It Solves")
para("Buying soft toys online is fragmented and generic. WonderFox gives a single branded, mobile-responsive storefront where a customer can browse a curated catalogue by category, search and filter products, sort, save items to a wishlist, pay by cash or online payment, and track order status - while the owner tracks every number from one dashboard.")

h2("1.5 Key Features (Implemented)")
bullets([
    "Customer: home page, shop-by-collection, product search/filter/sort + pagination, product details with reviews, add-to-cart, guest cart via localStorage, wishlist (guest + logged-in), quantity updates, COD + Razorpay checkout, order history with status.",
    "Auth: register, login, JWT bearer token stored in localStorage, protected customer routes.",
    "Admin: protected admin routes. Login, dashboard stats, product CRUD, category CRUD, order list/details + status update, user management with order summaries, blog CRUD, newsletter subscribers + send.",
    "Integrations: Cloudinary image upload, Razorpay payments with HMAC-SHA256 verification, nodemailer for newsletter + contact emails, SEO product meta tags.",
])

h2("1.6 Technology Stack (Summary)")
table(
    ["Tier", "Technology", "Where"],
    [
        ["Frontend", "React 19, Vite 8, Tailwind CSS 4", "storefront + admin"],
        ["UI / UX", "React Router 7, Framer Motion, Lucide React, React Hot Toast", "storefront + admin"],
        ["HTTP client", "Axios (with auth interceptors)", "both frontends"],
        ["Backend", "Node.js, Express 5", "REST API"],
        ["Database", "MongoDB (Atlas) + Mongoose 9 ODM", "all collections"],
        ["Auth", "jsonwebtoken (JWT), bcryptjs", "auth middleware + user model"],
        ["Files", "Multer + Cloudinary", "product / category / blog images"],
        ["Payments", "Razorpay (test keys)", "checkout"],
        ["Email", "nodemailer (SMTP)", "newsletter + contact"],
        ["Deploy", "Render, Vercel, Netlify configs", "frontend + API"],
    ],
    widths=[1.0, 2.3, 2.3],
)

h2("1.7 Why This Tech Stack Was Selected")
bullets([
    "React + Vite: fast component UI, instant dev server, ecosystem for routing, forms and animation.",
    "Tailwind CSS: utility-first styling keeps components responsive without a lot of custom CSS.",
    "Node.js + Express: lightweight, async request handling, huge middleware ecosystem (CORS, Multer, Morgan).",
    "MongoDB + Mongoose: flexible schemas, fast iteration, plus an ODM that adds validation and relationships.",
    "JWT + bcryptjs: stateless auth and safe password hashing.",
    "Razorpay + Cloudinary + Nodemailer: real-world integrations typical of production shops.",
    "Separate admin app keeps risky operations out of the public storefront.",
])

h2("1.8 Overall Request / Response / Data Flow")
code("""
 React storefront (browser)
      │ axios.get('/products?search=teddy&page=1')
      ▼
 Express Server (backend)
   app.use(cors) → app.use(express.json()) → morgan(log)
   → /api/products  (product.routes.js)
   → product.controller.getAllProducts
   → product.service.getAllProducts
   → Product.aggregate([ $match, $addFields, $facet ])
   → MongoDB  →  populate('category')
   → returns ApiResponse { success, statusCode, data, message }
      ▼
 Browser renders ProductCard list  (Collection.jsx)
""")

h2("1.9 The Three-Application Split")
bullet([
    "Storefront = customer-facing UI.",
    "API = the single source of truth / business logic layer.",
    "Admin = a trusted panel behind admin login.",
    "Both UI apps talk to the same backend API (same REST endpoints, different guards).",
])

# =====================================================================
# SECTION 2 - COMPLETE TECH STACK (part 1)
# =====================================================================
h1("2. Complete Tech Stack")

def tech(name, what, why, where, ans):
    h3(name)
    rich([("What it is: ", True), (what, False)])
    rich([("Why it is used in WonderFox: ", True), (why, False)])
    rich([("Where it is used: ", True), (where, False)])
    rich([("Interview-ready explanation: ", True), (ans, False)])

tech("React", "A component-based JavaScript library for building interactive UIs with reusable components and state.",
    "The storefront and admin are both built as React single-page apps. It lets the whole shop be split into small, testable components.",
    "frontend/src/pages/*.jsx, frontend/src/component/**; admin/src/pages and admin/src/components.",
    "React renders everything from components. Each page (like Collection.jsx) composes smaller components (ProductCard). State lives in useState hooks and context, and when it changes the component re-renders.")

tech("Vite", "A fast build tool and dev server for modern web apps.",
    "It starts both React apps instantly, gives hot reload while coding, and builds the production static bundle.",
    "frontend/vite.config.js, admin/vite.config.js; 'npm run dev' and 'npm run build'.",
    "Vite pre-bundles our React and Tailwind code into 'dist/' that hosts as a static site, and proxies '/api' to localhost in development.")

tech("Tailwind CSS", "A utility-first CSS framework where you style elements with small class names instead of custom stylesheets.",
    "It makes every card, button and page responsive quickly and keeps the design consistent (WonderFox orange).",
    "Every frontend and admin component, eg 'rounded-2xl bg-orange-500 py-3' in Login.jsx / ProductCard.jsx.",
    "Tailwind generates CSS from class names in the markup. Responsive layout uses prefixes like sm:, md:, lg: for screen sizes.")

tech("React Router", "A routing library that renders different pages for different URLs inside a React app.",
    "It creates multi-page behaviour without a page reload: /collection, /product/:id, /cart, /login, and admin routes.",
    "frontend/src/App.jsx, admin/src/App.jsx (BrowserRouter / Routes / Route).",
    "React Router maps a URL to a component. Dynamic parts are read with useParams (ProductDetails.jsx) and redirects use useNavigate.")

tech("Axios", "An HTTP client library that wraps fetch with interceptors, timeouts and error handling.",
    "It is the single network layer; both UIs call the backend, and interceptors attach the JWT automatically.",
    "frontend/src/api/axios.js, admin/src/api/axios.js; all services under */services/*.",
    "We create one axios instance with a baseURL. A request interceptor adds 'Bearer <token>'. A response interceptor catches 401 and redirects to /login.")

tech("Framer Motion", "An animation library for React that makes spring, fade and stagger animations simple.",
    "It adds scroll-reveal and drawer animations on the Collection page with one 'whileInView' prop.",
    "frontend/src/pages/Collection.jsx (product grid + filter drawer).",
    "Framer Motion wraps markup in motion.div. I set initial (hidden) and visible states; when the element scrolls into view it animates.")

tech("Lucide React", "A clean icon pack exported as React components.",
    "Used for icons (shopping bag, heart, eye, truck) across the shop UI.",
    "frontend/src/component/layout/Navbar.jsx, CartSummary.jsx, ProductInfo.jsx; admin too.",
    "Each icon is a React function imported from lucide-react; we size with width/height props and tint with className.")

tech("React Hot Toast", "A small toast-notification library for quick success / error feedback.",
    "Every form or action shows toast.success or toast.error instead of inline alert messages.",
    "AppProviders.jsx mounts <Toaster position='top-right'>; used in Login, Checkout, ProductCard and more.",
    "We call toast.success('...') after an action. The library renders a little notification box automatically and auto-dismisses.")

# =====================================================================
# SECTION 2 part 2 - backend tech
# =====================================================================
tech("Node.js", "A JavaScript runtime for running server-side JavaScript outside a browser.",
    "It hosts the Express API that bridges React and MongoDB.",
    "backend/src/server.js starts the server; 'npm run dev' via nodemon.",
    "Node handles many simultaneous HTTP requests on one thread using async/await, which matches MongoDB's async driver well.")

tech("Express.js", "A minimal web framework for Node that defines routes and middleware.",
    "It maps URLs to controllers, adds parsing for JSON and form data, cookie-parser, morgan logging and CORS.",
    "backend/src/app.js configures the app and mounts all /api route groups.",
    "Express builds a request pipeline; each middleware can alter the request or stop it. Our error middleware runs last.")

tech("MongoDB", "A NoSQL document database; each record is a JSON-like document.",
    "It stores users, products, carts, wishlists, orders, blogs and subscribers without a fixed schema.",
    "MongoDB Atlas cloud; backend/src/config/db.js connects with connectDB().",
    "MongoDB suits fast-evolving e-commerce data. Documents map to JSON, so the API and database match naturally.")

tech("Mongoose", "The MongoDB ODM (Object-Document Mapper) for Node.",
    "It defines schemas, adds validation and enums, handles references, and lets us populate related data.",
    "backend/src/models/*.js are all Mongoose schemas; services run queries with them.",
    "Mongoose models turn our JS schema into a collection, add validation and make queries like findById and populate easy.")

tech("JWT (jsonwebtoken)", "A signed, stateless token that proves a user's identity.",
    "After login the backend signs a token and returns it; protected routes verify it.",
    "backend/src/utils/generateToken.js; middleware/authMiddleware.js; stored in localStorage on the frontends.",
    "The token is a base64 header.payload.signature. protect middleware verifies the signature with JWT_SECRET and loads the user.")

tech("bcryptjs", "A password-hashing library implementing the slow bcrypt hash.",
    "Passwords are hashed before save so the raw value never touches the database.",
    "backend/src/models/User.js (pre-save hook) and the matchPassword method.",
    "bcrypt uses a per-password salt and is intentionally slow, which makes a leaked hash very hard to brute force.")

tech("REST API", "An API style where each resource is a URL endpoint and actions are HTTP verbs.",
    "It is the contract between the storefront, the admin panel and the backend.",
    "backend/src/routes/*.routes.js mounted under /api/...",
    "Products, carts and orders are resources; clients use GET / POST / PATCH / PUT / DELETE to read and change them.")

tech("Razorpay", "An online payment gateway (UPI, cards, wallets, net banking).",
    "Checkout offers COD and Razorpay. The backend creates and verifies orders using test keys.",
    "backend/src/config/razorpay.js and payment.service.js; frontend Checkout.jsx opens the Razorpay modal.",
    "The backend creates a Razorpay order with the amount in paise, the frontend opens the hosted checkout, then the backend verifies the HMAC-SHA256 signature before confirming the order.")

tech("Cloudinary", "Cloud image storage and CDN service.",
    "Product, category and blog images are uploaded to Cloudinary so the browser can serve them fast over a URL.",
    "backend/src/config/cloudinary.js, services/upload.service.js, routes/upload.routes.js.",
    "Multer reads the uploaded file into memory, it is base64-uploaded to Cloudinary, and the returned URL + public_id is stored on the document.")

tech("Multer", "A Node upload middleware that parses multipart/form-data.",
    "It reads uploaded image files into request memory (5 MB limit) before the Cloudinary upload.",
    "backend/src/middleware/uploadMiddleware.js.",
    "Multer parses the multipart body and exposes req.file.buffer, which upload.service then uploads to Cloudinary.")

tech("Nodemailer", "A Node library for sending email over SMTP.",
    "It powers the newsletter broadcast and the contact-form email using the Gmail SMTP account.",
    "backend/src/services/newsletter.service.js and contact.service.js.",
    "We build a nodemailer transporter from env vars, send mail per subscriber, and HTML-escape content to reduce XSS.")

tech("Other libraries", "cookie-parser (cookies), morgan (request logging), express-validator (installed), slugify (installed; Blog ships its own slugifier), react-icons, dotenv.",
    "They fill small gaps: logging, cookies, icons and reading env files.",
    "backend/package.json dependencies; frontend and admin package.json files.",
    "All are small, proven libraries; installing them avoids re-inventing logging, cookie parsing and env loading.")

# =====================================================================
# SECTION 3 - PROJECT ARCHITECTURE
# =====================================================================
h1("3. Project Architecture (Interview Language)")

h2("3.1 High-Level Picture")
code("""
      ┌──────────────┐     ┌──────────────┐
      │   frontend   │     │    admin     │
      │ (storefront) │     │ (dashboard)  │
      └──────┬───────┘     └──────┬───────┘
             │   axios + Bearer    │
             ▼                     ▼
      ┌─────────────────────────────────┐
      │     Express API (backend)        │
      │  routes -> controllers -> svc    │
      │  middleware protect/authorize    │
      └──────────────┬──────────────────┘
                     ▼  Mongoose
              MongoDB Atlas
""")

h2("3.2 Backend Layering")
table(
    ["Layer", "Role", "Real file"],
    [
        ["Routes", "Map URL + method to a controller, attach middleware", "routes/product.routes.js"],
        ["Middleware", "Reusable request filters (auth, role, upload, errors, CORS)", "middleware/authMiddleware.js"],
        ["Controllers", "Read request, call a service, send JSON", "controllers/product.controller.js"],
        ["Services", "Real business logic (validation, queries, transactions)", "services/order.service.js"],
        ["Models", "Mongoose schema + validation per collection", "models/*.js"],
        ["Utils/Helpers", "Shared tools (ApiError, token gen, price, stock)", "utils/*.js, helpers/*.js"],
        ["Config", "DB, env, Cloudinary, Razorpay setup", "config/*.js"],
    ],
    widths=[1.2, 3.3, 2.4],
)

h2("3.3 Request Pipeline (Middleware Order)")
code("""
 request -> cors -> express.json -> urlencoded -> cookieParser
        -> morgan(logger) -> /api/* routes
        -> (optional) protect + authorize
        -> controller -> service -> model query -> MongoDB
        -> ApiResponse JSON -> errorMiddleware (when an error is thrown)
""")

h2("3.4 Component / Route Structure")
bullets([
    "Storefront: / , /collection, /product/:id, /cart, /wishlist, /checkout, /orders, /order-success, /login, /register, /profile, /about, /contact, /blog, /blog/:slug, catch-all.",
    "Admin: / (login), /dashboard, /products + add/edit, /categories + add/edit, /orders + /orders/:id, /users + /users/:id, /blogs + add/edit, /newsletter.",
    "All guarded admin routes are wrapped in ProtectedRoute (checks localStorage token).",
])
para("Business logic lives in backend services, not in the React apps. Both frontends stay thin and the API is the single source of truth.")

h2("3.5 Database Layer")
bullet([
    "MongoDB Atlas (cloud) connected in config/db.js using MONGODB_URI.",
    "9 models: User, Product, Category, Cart, Wishlist, Order, Blog, NewsletterSubscriber, Review.",
    "Relationships use ObjectId references (Product.category, Order.user, Cart.user) populated with populate().",
    "Orders snapshot product data (name/image/price) so order history survives product edits.",
])

h2("3.6 Environment Variables (.env)")
table(
    ["Variable", "Purpose"],
    [
        ["PORT", "API port (5000)"],
        ["MONGODB_URI", "Atlas connection string"],
        ["JWT_SECRET / JWT_EXPIRE", "Token signing + lifetime (7d)"],
        ["CLIENT_URL / ADMIN_URL", "CORS allowlist origins"],
        ["CLOUDINARY_*", "Cloudinary API keys"],
        ["RAZORPAY_KEY_ID / _SECRET", "Razorpay keys (test)"],
        ["SMTP_HOST/PORT/USER/PASS/FROM_EMAIL", "nodemailer Gmail SMTP"],
        ["VITE_API_BASE_URL (both frontends)", "API base URL per environment"],
    ],
    widths=[2.5, 3.6],
)

# =====================================================================
# SECTION 4 - FOLDER STRUCTURE
# =====================================================================
h1("4. Folder Structure")
code("""
toy/
 ├─ frontend/                 React storefront (customer)
 │   └─ src/
 │       ├─ api/axios.js          Axios instance + JWT interceptors
 │       ├─ component/           Reusable UI: cart/, checkout/, common/,
 │       │                        home/, layout/, product/, ui/, about/
 │       ├─ context/             CartContext (guest cart in localStorage)
 │       ├─ data/                static demo data (categories, testimonials)
 │       ├─ hooks/               useCart / useWishlist / useAuth hooks
 │       ├─ layouts/MainLayout   Navbar + <Outlet/> + Footer
 │       ├─ pages/               Home, Collection, ProductDetails, Cart,
 │       │                        Wishlist, Login, Register, Profile,
 │       │                        Checkout, MyOrders, OrderSuccess, Blog,
 │       │                        BlogDetail, About, Contact, NotFound
 │       ├─ providers/           AppProviders (Toaster + CartProvider)
 │       ├─ services/            API wrappers (auth, cart, order...)
 │       └─ utils/authGuard.js   login-redirect helpers
│
├── backend/
│   └─ src/
│       ├─ app.js / server.js    Express app + entrypoint
│       ├─ config/               env, db, cloudinary, razorpay
│       ├─ constants/
│       ├─ controllers/          auth, product, cart, order, payment ...
│       ├─ helpers/              order total, stock update, snapshots
│       ├─ middleware/           authMiddleware, adminMiddleware, upload, error
│       ├─ models/               User, Product, Category, Cart, Wishlist,
│       │                        Order, Blog, NewsletterSubscriber, Review
│       ├─ routes/               all /api route groups
│       ├─ services/             business-logic layer
│       ├─ utils/                ApiError, ApiResponse, generateToken ...
│       └─ validators/
│   └─ uploads/                  local upload placeholder
│
├── admin/                       React dashboard (owner)
│   └─ src/
│       ├─ api/axios.js
│       ├─ components/          common/, dashboard/, layout/, products/, blog/
│       ├─ context/ hooks/ utils/
│       ├─ layouts/AdminLayout   Sidebar + Navbar
│       ├─ pages/                auth/, dashboard/, products/, categories/,
│       │                        orders/, users/, blogs/, newsletter/
│       ├─ routes/ProtectedRoute
│       └─ services/             API wrappers
│
├── netlify.toml  vercel.json  render.yaml   static deploy configs
├── FEATURE_SUMMARY.md  ADMIN_CONFIG_FIX.md  RAZORPAY_INTEGRATION.md
└── generate_docx.py  (older generator) + build_docx.py (this file)
""")

h2("4.1 Why this structure matters (interview answer)")
para("The repo is split into three installable/runable apps plus a shared mental model: the storefront and dashboard are both thin React clients, while the backend owns all business rules. This gives clean separation of concerns, lets two teams scale independently, and keeps sensitive admin logic away from shoppers. Services/ and controllers/ separate routing from business logic, so route files stay tiny and queries stay reusable.")

# =====================================================================
# SECTION 5 - COMPLETE USER FLOW
# =====================================================================
h1("5. Complete User Flow")

h2("5.1 Customer Flow (implemented)")

h3("1. Opening the website")
para("User lands on '/' (Home). MainLayout renders Navbar + Footer. Home pulls static content and shows featured categories (data/*.js).")

h3("2. Browsing products")
para("From the navbar the user opens '/collection'. Collection.jsx calls getCategories() and getProducts() in parallel and renders ProductCard grids.")

h3("3. Searching / filtering / sorting")
para("The navbar search navigates to '/collection?search=teddy'. The filter drawer lets the user set category, minPrice, maxPrice, sort. Every change updates URL search params (useSearchParams), which re-fires the product fetch. Backend applies $regex search, price range on sellingPrice, and sorts by allowed fields.")

h3("4. Product details")
para("'/product/:id' (ProductDetails.jsx) fetches GET /products/:id and GET /products (for related items), then renders ProductGallery, ProductInfo (price, discount %, quantity selector, add-to-cart, buy-now), ProductTabs (description/reviews) and RelatedProducts. It also injects SEO meta + JSON-LD product schema.")

h3("5. Adding to cart")
para("Clicking 'Add To Cart' calls useCart().addItem(product._id, quantity). If a token exists it calls POST /cart; otherwise it normalizes a product snapshot into localStorage key 'guestCart' and dispatches a 'cart:updated' event.")

h3("6. Updating cart")
para("Cart.jsx lists CartItem components with +/- quantity and Remove buttons. Quantity updates call PATCH /cart/:productId (backend refuses quantity above stock, and the increase button is disabled when qty >= stock).")

h3("7. Login / signup")
para("'/login' and '/register' are plain controlled forms. Login stores token + user in localStorage, merges guest cart by calling addToCart per guest item, then navigates to the ?redirect= URL or '/'. Register simply creates the account and navigates home.")

h3("8. Checkout")
para("CartSummary 'Proceed to Checkout' requires login (redirects to /login?redirect=/checkout if not). Checkout.jsx validates the Indian address (10-digit phone starting 6-9, 6-digit pincode), lets the user pick COD or Razorpay, and shows DeliveryAddress + OrderSummary components.")

h3("9. Order creation (COD)")
para("Place Order → POST /orders { shippingAddress }. Backend starts a MongoDB transaction: verifies cart, builds order snapshot, creates Order with paymentMethod COD, decrements product stock via bulkWrite, clears cart, commits.")

h3("10. Payment (Razorpay) - implemented")
para("Place Order → POST /orders/razorpay returns { orderId, amount, currency, key }. The frontend opens new window.Razorpay(options). On success the handler POSTs /orders/verify-payment with razorpayOrderId/paymentId/signature + shippingAddress. Backend verifies the HMAC signature, then creates the order (paymentStatus PAID, orderStatus CONFIRMED) in a transaction.")

h3("11. Order history")
para("'/orders' (MyOrders.jsx) calls GET /orders and renders order numbers, badges per orderStatus, item snapshots and totals.")

h3("12. Logout")
para("Profile.jsx clears localStorage token + user and navigates to /login. There is no backend logout endpoint (JWT is stateless).")

h2("5.2 Admin Flow (implemented)")
table(
    ["Step", "How it works", "Files"],
    [
        ["Admin login", "Admin logs in via /auth/login (same endpoint as customers). Admin role is a field on the User doc, seeded manually.", "admin/pages/auth/Login.jsx"],
        ["Route guard", "ProtectedRoute redirects to '/' if localStorage has no token.", "admin/routes/ProtectedRoute.jsx"],
        ["Dashboard", "GET /dashboard returns totals (products, categories, orders, revenue) + recent products + low-stock products (stock <= 5).", "admin/pages/dashboard/Dashboard.jsx"],
        ["Product management", "Products list, Add Product, Edit Product. Uses POST/PUT/DELETE /products (protected) + image upload via /upload.", "admin/pages/products/*"],
        ["Category management", "Categories list/add/edit with optional image upload (multipart form-data).", "admin/pages/categories/*"],
        ["Order management", "Orders list via GET /orders/admin, details via /orders/admin/:id, status update via PATCH /orders/admin/:id/status.", "admin/pages/orders/*"],
        ["User management", "Users list + details with number of orders, total spent and recent orders.", "admin/pages/users/*"],
        ["Blogs", "Blog list/add/edit/delete with slug generation.", "admin/pages/blogs/*"],
        ["Newsletter", "View subscribers, deactivate, and send a newsletter email to all active subscribers.", "admin/pages/newsletter/*"],
    ],
    widths=[1.4, 3.6, 1.9],
)
para("Note: backend product create/update endpoints use the protect middleware (any logged-in user) - the stronger admin authorize guard is applied to orders, users, blogs, newsletter and dashboard routes.")

# ====MORE====