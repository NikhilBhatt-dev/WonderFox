#!/usr/bin/env python
"""Generate 'WonderFox E-Commerce Store - Complete Interview Preparation Guide' DOCX.

Content is derived strictly from the actual WonderFox codebase (toy/backend,
toy/frontend, toy/admin). Features that are NOT implemented in the code are
explicitly marked as "Planned / Future Enhancement".
"""

import os
import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(BASE_DIR, "WonderFox_Interview_Guide.docx")


# ----------------------------------------------------------------------------
# Styling helpers
# ----------------------------------------------------------------------------

_bookmark_counter = {"n": 0}


def _next_bookmark_id():
    _bookmark_counter["n"] += 1
    return _bookmark_counter["n"]


def _sanitize_bookmark(name):
    """Word bookmark names may not contain spaces; keep alnum + underscore."""
    return "".join(c if (c.isalnum() or c == "_") else "_" for c in str(name))


def add_bookmark(paragraph, name):
    """Attach a Word bookmark to a paragraph so it is jumpable from the TOC."""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(_next_bookmark_id()))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(_bookmark_counter["n"]))
    p = paragraph._p
    p.insert(0, start)
    p.append(end)


def add_field(paragraph, field_code, fallback=""):
    """Insert a Word field (e.g. PAGE, NUMPAGES) into a paragraph."""
    run = paragraph.add_run()
    r = run._r
    b1 = OxmlElement("w:fldChar")
    b1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = field_code
    b2 = OxmlElement("w:fldChar")
    b2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = fallback
    b3 = OxmlElement("w:fldChar")
    b3.set(qn("w:fldCharType"), "end")
    for el in (b1, instr, b2, t, b3):
        r.append(el)


# A tiny markdown-lite parser: supports **bold** and `code`.
_TOKEN_RE = re.compile(r"(\*\*|` )")
_BOLD = "**"
_CODE = "`"


def _add_rich_runs(paragraph, text, size=11, bold_default=False, code_default=False):
    """Render `**bold**` and `` `code` `` into styled runs on a paragraph."""
    i = 0
    n = len(text)
    bold = bold_default
    code = code_default
    buf = []

    def flush():
        if not buf:
            return
        txt = "".join(buf)
        run = paragraph.add_run(txt)
        run.font.name = "Consolas" if code else "Calibri"
        run.font.size = Pt(size)
        run.bold = bold
        if code:
            run.font.name = "Consolas"

    while i < n:
        if text[i:i + 2] == _BOLD:
            flush()
            buf.clear()
            bold = not bold
            i += 2
            continue
        if text[i] == _CODE:
            flush()
            buf.clear()
            code = not code
            i += 1
            continue
        buf.append(text[i])
        i += 1
    flush()


def para(doc, text="", size=11, align="left", spacing=0, keep_next=False,
         keep_together=False):
    """Add a paragraph, parsing `**bold**` and `` `code` `` inline."""
    par = doc.add_paragraph()
    _add_rich_runs(par, text, size=size)
    for r in par.runs:
        r.font.name = r.font.name  # preserve Consolas/Calibri set above
    pf = par.paragraph_format
    pf.space_before = Pt(spacing)
    pf.space_after = Pt(spacing)
    pf.line_spacing = 1.15
    if align == "center":
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "justify":
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if keep_next:
        pf.keep_with_next = True
    if keep_together:
        pf.keep_together = True
    return par


def heading(doc, text, level, bookmark=None, keep_next=True):
    par = doc.add_paragraph()
    if bookmark:
        add_bookmark(par, bookmark)
    run = par.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(14 if level == 1 else 12)
    run.bold = True if level <= 2 else False
    run.font.color.rgb = RGBColor(26, 31, 55)
    pf = par.paragraph_format
    pf.space_before = Pt(12) if level == 1 else Pt(6)
    pf.space_after = Pt(4)
    pf.keep_with_next = keep_next
    par.style = "Heading {}".format(level)
    return par


def subheading(doc, text, bookmark=None):
    return heading(doc, text, 2, bookmark=bookmark)


def h3(doc, text, bookmark=None):
    return heading(doc, text, 3, bookmark=bookmark, keep_next=True)


def h4(doc, text):
    return heading(doc, text, 4, keep_next=False)


def bullet(doc, items, size=11):
    for it in items:
        par = doc.add_paragraph(style="List Bullet")
        _add_rich_runs(par, it, size=size)
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.line_spacing = 1.15
    return bullet


def numbered(doc, items, start=1, size=11):
    for it in items:
        par = doc.add_paragraph(style="List Number")
        _add_rich_runs(par, it, size=size)
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.line_spacing = 1.15


def note(doc, text, label="Note", color=(130, 80, 0)):
    par = doc.add_paragraph()
    run = par.add_run("[{}] ".format(label))
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = RGBColor(*[c for c in color])
    run2_par = par.add_run()
    _add_rich_runs_to_run(run2_par.add_run(""), text)  # placeholder, replaced below
    # Replace: clear the placeholder run and render rich text
    run2 = par.add_run()
    run2.font.name = "Calibri"
    run2.font.size = Pt(10)
    _add_rich_runs_in_run(par, text, run2)
    par.paragraph_format.left_indent = Pt(6)
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after = Pt(4)
    return par


def _add_rich_runs_in_run(paragraph, text, run):
    """Append styled runs to an existing paragraph (helper for note labels)."""
    # Simplistic: reuse rich logic by adding directly to paragraph.
    tmp = paragraph
    # We already added `run`; remove it then use _add_rich_runs on the paragraph
    # and discard the empty run.
    pass


def planned(doc, text):
    return note(doc, text, label="Planned / Future Enhancement",
                color=(90, 90, 90))


def sub_note(doc, text):
    par = doc.add_paragraph()
    run = par.add_run()
    _add_rich_runs_to_run(run, text)
    par.paragraph_format.left_indent = Pt(6)
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after = Pt(4)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    return par


def _add_rich_runs_to_run(run, text, size=10):
    """Render `**bold**`/`` `code` `` into runs inside a paragraph, appended."""
    # We cannot set font per-run easily via run; create separate runs on run's paragraph.
    paragraph = run._p.getparent() if False else run._r.getparent()
    # run._r is the <w:r>; its parent is <w:p>. We add sibling <w:r> elements.
    # Simpler: just add the whole text to the given run (plain).
    run.text = text
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    return run


def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for r in hdr_cells[i].paragraphs:
            for rr in r.runs:
                rr.bold = True
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for r in cells[i].paragraphs:
                for rr in r.runs:
                    rr.font.name = "Calibri"
                    rr.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in t.columns[i].cells:
                cell.width = w
    return t


def code_block(doc, code, size=8.5):
    par = doc.add_paragraph()
    run = par.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(size)
    pf = par.paragraph_format
    pf.left_indent = Pt(12)
    pf.right_indent = Pt(12)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.0
    # Light gray background
    pPr = par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    return par


# ----------------------------------------------------------------------------
# TOC collection
# ----------------------------------------------------------------------------

TOC = []  # list of (level, title, bookmark)


def reg(title, level, bookmark=None):
    bkm = _sanitize_bookmark(bookmark) if bookmark else _sanitize_bookmark(title)
    TOC.append((level, title, bkm))
    return bkm


def build_toc(doc, marker_par):
    """Insert a clickable TOC right before `marker_par`, then remove the marker."""
    made = []
    for level, title, bookmark in TOC:
        par = doc.add_paragraph()
        par.paragraph_format.left_indent = Pt(14 + (level - 1) * 18)
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.line_spacing = 1.15
        hl = OxmlElement("w:hyperlink")
        hl.set(qn("w:anchor"), bookmark)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "4472C4")
        rPr.append(color)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = title
        if title.startswith(" "):
            t.set(qn("xml:space"), "preserve")
        r.append(t)
        hl.append(r)
        par._p.append(hl)
        made.append(par)
    # Move each TOC paragraph to just before the marker.
    for par in made:
        marker_par._p.addprevious(par._p)
    # Drop the marker.
    parent = marker_par._p.getparent()
    if parent is not None:
        parent.remove(marker_par._p)


def footer_with_page_numbers(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    footer = section.footer
    fpar = footer.paragraphs[0]
    fpar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fpar.add_run()
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    add_field(fpar, "PAGE", "2")
    run.add_text(" of ")
    add_field(fpar, "NUMPAGES", "30")


# ----------------------------------------------------------------------------
# Document build
# ----------------------------------------------------------------------------

doc = Document()

# Default style
_style = doc.styles["Normal"]
_font = _style.font
_font.name = "Calibri"
_font.size = Pt(11)

# Section: page margins + footer with page numbers (title page first page excluded)
section = doc.sections[0]
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)


def title_page():
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("WonderFox E-Commerce Store")
    r.font.name = "Calibri"
    r.font.size = Pt(32)
    r.bold = True
    r.font.color.rgb = RGBColor(26, 31, 55)
    r = title.add_run("\n")
    r = title.add_run("Complete Interview Preparation Guide")
    r.font.name = "Calibri"
    r.font.size = Pt(20)
    r.bold = True
    r.font.color.rgb = RGBColor(26, 31, 55)
    doc.add_paragraph()
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("A technical interview reference based entirely on the\nWonderFox codebase (backend / frontend / admin).")
    r.font.name = "Calibri"
    r.font.size = Pt(13)
    doc.add_paragraph()
    doc.add_paragraph()
    footer_with_page_numbers(doc)
    # Marker that the TOC will replace.
    marker = doc.add_paragraph("\x00TOC_MARKER")
    return marker


TOC_MARKER = title_page()

# ---- TOC heading ----
heading(doc, "Table of Contents", 1, bookmark="toc")
sub_note(doc, "Internal section titles below are clickable bookmarks that jump to each part of the guide.")

# ===================== SECTION 1 =====================
b1 = reg("1. Project Overview", 1)
heading(doc, "1. Project Overview", 1, bookmark=b1)
para(doc, "WonderFox is a full-stack e-commerce store built for selling premium soft toys and kids' toys online. The codebase is split into three independently runnable projects:")
bullet(doc, [
    "**toy/backend** - Node.js + Express + MongoDB API",
    "**toy/frontend** - React + Vite customer storefront",
    "**toy/admin** - React + Vite admin dashboard",
])
note(doc, "Throughout the guide, file paths such as `toy/frontend/src/...` refer to the monorepo layout under the `toy/` workspace folder.")
subheading(doc, "1.1 What is WonderFox?", bookmark="1_1")
para(doc, "**WonderFox** is a toy e-commerce application. The customer-facing store (frontend) lets shoppers browse toys, search and filter products, manage a cart, maintain a wishlist, checkout with Cash on Delivery or Razorpay, track orders, and read a blog. The admin panel lets an administrator manage products, categories, orders, users, blogs and newsletters. The Express backend exposes a REST API and handles authentication, payments, image uploads (Cloudinary), email (nodemailer), and MongoDB persistence.")
subheading(doc, "1.2 Purpose of the project", bookmark="1_2")
para(doc, "The project demonstrates a complete MERN-style e-commerce stack: product catalog, cart/wishlist, JWT authentication, role-based authorization (customer vs admin), order placement with stock management, Razorpay payment integration, an admin CMS, a blog, and a newsletter system. It is designed as an interview/portfolio project that a developer can explain end-to-end.")
subheading(doc, "1.3 Target users", bookmark="1_3")
bullet(doc, [
    "**Customers** - browse, search, buy toys, manage account, view order history.",
    "**Administrators** - manage catalog, orders, users, content, and subscribers.",
])
subheading(doc, "1.4 Business problem solved", bookmark="1_4")
para(doc, "Parents want a trustworthy single place to discover safe, premium kids' toys with a frictionless buying journey (guest cart, simple checkout, multiple payment options) while merchants/admins need a lightweight, self-service panel to run the store without custom engineering for every catalog change.")
subheading(doc, "1.5 Key features", bookmark="1_5")
bullet(doc, [
    "Product listing with search, category, price, sort and pagination (backend aggregation).",
    "Product detail page with gallery, related products, reviews, SEO metadata and JSON-LD structured data.",
    "Guest cart (localStorage) that merges into the authenticated cart after login.",
    "Guest wishlist (localStorage IDs) plus authenticated wishlist backed by MongoDB.",
    "JWT authentication (register/login/me) with token stored in localStorage and role-based authorization.",
    "Cart add/update/remove/clear with live stock validation.",
    "Checkout with address form validation, COD ordering, and Razorpay (order create + signature verification).",
    "Order history for customers; full order status flow with status history.",
    "Admin: product/category/blog/order/user CRUD and dashboard statistics.",
    "Newsletter subscribe (public) + admin send/deactivate.",
    "Contact form that emails the merchant via nodemailer.",
    "Image upload to Cloudinary (products/categories/blog images).",
])
subheading(doc, "1.6 Technology stack (summary)", bookmark="1_6")
bullet(doc, [
    "**Backend:** Node.js, Express 5, MongoDB (Mongoose 9), bcryptjs, jsonwebtoken, cloudinary, multer, razorpay, nodemailer, slugify, cors, cookie-parser, morgan, dotenv",
    "**Frontend:** React 19, Vite 8, Tailwind CSS 4 (via @tailwindcss/vite), React Router DOM 7, Axios, Framer Motion, Lucide React, React Hot Toast, react-icons",
    "**Admin:** React 19, Vite 8, Tailwind CSS 4, React Router DOM 7, Axios, Lucide React, React Hot Toast",
])
subheading(doc, "1.7 Why this tech stack was selected", bookmark="1_7")
bullet(doc, [
    "MERN-family stack so every layer speaks JavaScript, reducing context switching.",
    "Express 5 + Mongoose gives a thin, fast REST API with schema validation close to the data.",
    "Vite gives near-instant dev/compile for both React apps.",
    "Tailwind CSS 4 JIT means a consistent design system (brand colors defined in CSS custom properties) with tiny CSS output.",
    "Cloudinary removes the need to host/resize images ourselves.",
    "Razorpay is India's dominant payments SDK with well-documented signature verification.",
])
subheading(doc, "1.8 Frontend / backend / database / API architecture", bookmark="1_8")
para(doc, "**Architecture:** the frontend and admin are static SPAs (built to `dist/`, served as static sites with SPA rewrites), the backend is a stateful Express API, and MongoDB is the single data store. The API is REST/JSON consumed by Axios. Cross-origin requests are allowed for `CLIENT_URL` and `ADMIN_URL` plus localhost dev origins, with `credentials: true`.")
numbered(doc, [
    "Browser (React) calls a service (`/services/*.service.js`) which uses the shared `api` Axios instance.",
    "Axios `request` interceptor injects the `Authorization: Bearer <token>` header from `localStorage`.",
    "Request hits Express -> `app.js` CORS/JSON/cookie/morgan middleware -> matched route.",
    "Route calls `protect` (JWT verify) and optionally `authorize('admin')`.",
    "Controller delegates to a service, which validates input, queries/updates Mongoose models, and returns an `ApiResponse` (or throws `ApiError`).",
    "The global `errorMiddleware` (registered last) catches the error and sends `{ success, message }`.",
    "Axios `response` interceptor detects 401, clears stale auth, and redirects to `/login?redirect=...`.",
    "React re-renders from `useState`/`useContext` state returned by the service.",
])
note(doc, "Frontend stores the JWT in **localStorage** (not an HttpOnly cookie). The backend has `cookie-parser` installed but **cookies are not used for auth** - the token is sent as a Bearer header. This is a deliberate trade-off (stateless JWT) and is noted under Security limitations.")


# ===================== SECTION 2 =====================
b = reg("2. Complete Tech Stack", 1)
heading(doc, "2. Complete Tech Stack", 1, bookmark=b)
para(doc, "Only technologies actually present in `package.json` / imports are listed below. Each entry states: what it is, where it is used, and why WonderFox uses it.")

TECH = [
    ("React", "A JavaScript library for building component-based user interfaces.",
     "Frontend (`toy/frontend/src`) and Admin (`toy/admin/src`).",
     "`useState` for component state, `useEffect` for side-effects, component composition for pages and reusable UI (e.g. ProductCard, CartItem, ProductTabs). Built with React 19."),
    ("Vite", "Next-generation build tool and dev server.",
     "`frontend/vite.config.js` and `admin/vite.config.js`; scripts `vite`, `vite build`, `vite preview`.",
     "Replaces webpack/CRA - fast HMR and bundling; both apps run independently on separate dev ports."),
    ("Tailwind CSS 4", "Utility-first CSS framework; v4 ships via the `@tailwindcss/vite` plugin (no tailwind.config.js needed).",
     "`src/index.css` (`@import \"tailwindcss\"`) and `src/App.css`. Custom brand colors are defined as CSS custom properties in `frontend/src/index.css` (`@theme`).",
     "Gives a consistent design system (blue primary, orange secondary, green accent) with utility classes everywhere and tiny CSS output."),
    ("React Router DOM 7", "Declarative client-side routing for React.",
     "`frontend/src/App.jsx` (`BrowserRouter` / `Routes`) and admin `App.jsx`.",
     "Enables SPA navigation and nested layouts (`MainLayout` / `Outlet`, `AdminLayout`)."),
    ("Axios", "HTTP client with interceptors.",
     "`frontend/src/api/axios.js` (frontend) and `admin/src/api/axios.js` (admin). All services import this `api`.",
     "Centralizes base URL (`VITE_API_BASE_URL`), attaches the Bearer token on every request, and automatically handles 401 by clearing auth and redirecting."),
    ("Framer Motion", "Production-ready animation library.",
     "`frontend/src/component/home/`, `pages/Collection.jsx` (`motion.div`, `AnimatePresence`).",
     "Staggered reveal animations, hover lifts and the mobile filter drawer slide-in - without writing CSS keyframes."),
    ("Lucide React", "Icon library (SVG components).",
     "Frontend (`Navbar`, `ProductInfo`, `CartItem`, `CheckoutStepper`, etc.) and Admin (`Sidebar`, `Navbar`).",
     "Consistent, accessible, tree-shaken icons (ShoppingCart, Heart, CheckCircle2)."),
    ("React Hot Toast", "Toast notification system.",
     "`providers/AppProviders.jsx` renders `<Toaster>`; calls appear in Login, CartItem, Checkout, Contact, Blog, ProductTabs.",
     "Immediate user feedback for async actions (add to cart, order placement, errors)."),
    ("react-icons", "Additional icon sets (Font Awesome 6 brands).",
     "`frontend/src/component/layout/Footer.jsx` (`FaFacebookF`, `FaInstagram`, `FaXTwitter`, `FaYoutube`).",
     "Social-brand icons not available in Lucide."),
    ("Node.js", "JavaScript runtime (backend).",
     "`toy/backend/src/server.js`; all backend code runs on Node.",
     "Lets us run the same language (JS) on server and client."),
    ("Express.js 5", "Web application framework.",
     "`toy/backend/src/app.js` - defines middleware, mounts all routers, health-check route, and the error handler (last).",
     "Minimal, unopinionated framework that maps URLs to controller functions."),
    ("MongoDB", "Document database (NoSQL).",
     "Connected in `toy/backend/src/config/db.js` via `mongoose.connect(process.env.MONGODB_URI)`.",
     "Flexible schema for products/orders while Mongoose enforces structure."),
    ("Mongoose 9", "MongoDB ODM: schemas, validation, population and aggregation.",
     "`toy/backend/src/models/*.js` (9 models) and `toy/backend/src/services/*.js`.",
     "Provides `pre('save')` hooks, ref/populate, `$inc` stock updates, and the aggregation pipeline used for product listing."),
    ("bcryptjs", "Pure-JS bcrypt (no native binding).",
     "`toy/backend/src/models/User.js` (`pre('save')` hashes password with 10 rounds) and `matchPassword`.",
     "Hashes passwords so the DB never stores plaintext credentials."),
    ("jsonwebtoken (JWT)", "Token creation/verification.",
     "`toy/backend/src/utils/generateToken.js` signs `{ id }` with `JWT_SECRET` / `JWT_EXPIRE`; `middleware/authMiddleware.js` verifies it.",
     "Stateless authentication: the token carries the user id and expiry."),
    ("Express CORS", "Cross-Origin Resource Sharing middleware.",
     "`toy/backend/src/app.js` - `cors({ origin, credentials: true })` with a dynamic origin whitelist.",
     "Lets the separate frontend/admin origins (and localhost dev) call the API with credentials."),
    ("cookie-parser", "Parses `Cookie` / `Set-Cookie` headers.",
     "Imported and `app.use(cookieParser())` in `app.js`.",
     "**Installed but unused for auth** - token goes in the `Authorization` Bearer header; cookies are not used for auth."),
    ("Morgan", "HTTP request logger.",
     "`toy/backend/src/app.js` - `morgan('dev')`.",
     "Logs every request (method, URL, status) to the console for debugging."),
    ("dotenv", "Loads `.env` into `process.env`.",
     "`toy/backend/src/config/env.js` calls `dotenv.config()` pointing at the repo `.env`.",
     "Keeps secrets (DB URL, JWT secret, Razorpay keys, SMTP creds) out of source code."),
    ("Multer", "Multipart/form-data (file upload) middleware.",
     "`toy/backend/src/middleware/uploadMiddleware.js` (memory storage, 5 MB limit); used in upload & category routes.",
     "Streams image uploads into memory so they can be forwarded straight to Cloudinary."),
    ("Cloudinary", "Cloud image management.",
     "`toy/backend/src/config/cloudinary.js` + `toy/backend/src/services/upload.service.js`.",
     "Stores product/category/blog images under the `wonderfox` folder and returns a secure `url` + `public_id`."),
    ("Razorpay", "Indian payments gateway SDK.",
     "`toy/backend/src/config/razorpay.js` + `toy/backend/src/services/payment.service.js`.",
     "Creates Razorpay orders (amount in paise) and verifies HMAC signatures server-side."),
    ("Nodemailer", "Email transport.",
     "`services/newsletter.service.js` (SMTP transport) and `services/contact.service.js`.",
     "Sends the contact-form email and bulk newsletter emails to active subscribers."),
    ("Slugify", "URL slug generation.",
     "`toy/backend/src/services/blog.service.js` (`slugify` helper + unique-slug loop).",
     "Creates SEO-friendly blog URLs."),
    ("express-validator", "Declarative request validation (declared in package.json).",
     "Listed in `toy/backend/package.json`.",
     "**Declared but unused** - WonderFox validates manually in utils/helpers (`validateQuantity`, `validateObjectId`, `validateOrderItems`) and services (`validateBlogPayload`). No import of `check`/`express-validator` exists in the code."),
    ("Nodemon", "Auto-restarts the dev server.",
     "`toy/backend/package.json` dev dependency + `npm run dev`.",
     "Rebuilds/restarts the backend on file changes during development."),
    ("ESLint (flat config)", "Linting.",
     "`frontend/eslint.config.js`.",
     "Keeps JSX and React-hooks rules consistent in the frontend."),
]
for name, what, where, explain in TECH:
    subheading(doc, name, bookmark="tech_" + name.replace(" ", "_"))
    para(doc, "**What:** " + what)
    para(doc, "**Used in:** " + where)
    para(doc, "**Why / how it's used:** " + explain)


# ===================== SECTION 3 =====================
b = reg("3. Backend Architecture", 1)
heading(doc, "3. Backend Architecture", 1, bookmark=b)
para(doc, "The backend (`toy/backend`, package name `backend`, type: module) follows a clean layered pattern: **Route -> Controller -> Service -> Model**. Controllers are intentionally thin (they unwrap the request and delegate), because every piece of business logic lives in a service that returns a uniform `ApiResponse` (or throws `ApiError`).")

subheading(doc, "3.1 Entry point & configuration", bookmark="3_1")
para(doc, "`server.js` is the only entry point. It loads dotenv from `../../.env`, connects to MongoDB, then starts the Express app:")
code_block(doc, """import "./config/env.js";
import app from "./app.js";
import connectDB from "./config/db.js";

const PORT = process.env.PORT || 5000;
const startServer = async () => {
  try {
    await connectDB();
    app.listen(PORT, () => console.log(`Server running on http://localhost:${PORT}`));
  } catch (error) {
    console.error(error);
  }
};
startServer();""")
note(doc, "`connectDB` uses `mongoose.connect(process.env.MONGODB_URI)` and crashes the process (`process.exit(1)`) on failure - simple but acceptable for a deploy-once service.")

subheading(doc, "3.2 Middleware stack (app.js)", bookmark="3_2")
numbered(doc, [
    "`cors({ origin, credentials: true })` - dynamic origin whitelist built from `CLIENT_URL`, `ADMIN_URL` and `http://localhost:5173`; an unknown origin is rejected with `CORS blocked for origin: <origin>`.",
    "`express.json()` - parses JSON request bodies.",
    "`express.urlencoded({ extended: true })` - parses URL-encoded bodies.",
    "`cookieParser()` - present but cookies are NOT used for auth (JWT is in the header).",
    "`morgan('dev')` - HTTP request logging.",
    "All routers mounted under `/api/*`.",
    "`errorMiddleware` registered **last** so it catches every thrown `ApiError`.",
])
note(doc, "Health-check route `GET /` returns `{ success, message: \"WonderFox Backend Running\" }`.")

subheading(doc, "3.3 Routing & the Route->Controller->Service->Model pipeline", bookmark="3_3")
para(doc, "There are 12 routers mounted in `app.js`:")
bullet(doc, [
    "auth -> `/api/auth`",
    "products -> `/api/products`",
    "categories -> `/api/categories`",
    "upload -> `/api/upload`",
    "cart -> `/api/cart`",
    "wishlist -> `/api/wishlist`",
    "orders -> `/api/orders`",
    "dashboard -> `/api/dashboard`",
    "users -> `/api/users`",
    "blogs -> `/api/blogs`",
    "newsletter -> `/api/newsletter`",
    "contact -> `/api/contact`",
])
note(doc, "Controllers wrap handlers in `asyncHandler` (which `Promise.resolve`s the handler and forwards rejections to `next`), do almost nothing except delegate to a service, and send the returned `ApiResponse`/`ApiError` via `sendResponse`. Example from `cart.controller.js`:")
code_block(doc, """export const addToCart = asyncHandler(async (req, res) => {
  const { productId, quantity } = req.body;
  const response = await cartService.addToCart(req.user._id, productId, quantity);
  return sendResponse(res, response);
});""")

subheading(doc, "3.4 Authentication & authorization", bookmark="3_4")
para(doc, "**JWT (Bearer auth).** `generateToken(id)` signs `{ id }` with `JWT_SECRET` for `JWT_EXPIRE`. The frontend stores the token in `localStorage` and the Axios request interceptor attaches `Authorization: Bearer <token>` to every call. The `protect` middleware extracts the token, verifies it, looks up the user (`User.findById(decoded.id).select('-password')`) and sets `req.user`. The `authorize(...roles)` middleware (role normalized + lower-cased) enforces admin-only routes using `protect` first, then role checks -> 401 if no user, 403 if wrong role.")

subheading(doc, "3.5 Error handling & the uniform response envelope", bookmark="3_5")
para(doc, "Two custom types standardize every response:")
bullet(doc, [
    "**`ApiError`** (utils/ApiError.js) - extends `Error`; carries `statusCode` and `success = false`. Thrown anywhere in services.",
    "**`ApiResponse`** (utils/ApiResponse.js) - envelope `{ success, statusCode, message, data }` returned by services.",
])
note(doc, "`errorMiddleware` is the final middleware. It logs the stack, then sends `{ success: false, message }` with `err.statusCode || 500`. Because it is registered last, it also catches synchronous `throw` inside `asyncHandler` (via `next(err)`).")

subheading(doc, "3.6 Validation strategy (manual, not express-validator)", bookmark="3_6")
para(doc, "Despite `express-validator` being a declared dependency, validation is performed manually:")
bullet(doc, [
    "`validateQuantity(qty)` - integer >= 1.",
    "`validateObjectId(id)` - `mongoose.Types.ObjectId.isValid`.",
    "`getValidProduct(id)` - ObjectId check + exists + `isActive` true.",
    "`validateOrderItems(cart, session)` - non-empty cart, valid ObjectIds, `isActive` products, integer qty >= 1, `product.stock >= qty`.",
    "`validateBlogPayload` (in blog.service.js) - required fields, cover image on create, enum checks.",
])
planned(doc, "Migrate to `express-validator` chain middleware (or Zod) on routes for request-shape validation; the dependency is already present.")

subheading(doc, "3.7 Transactions & data integrity", bookmark="3_7")
para(doc, "The two money-mutating flows - COD orders and Razorpay payment verification - wrap everything in a MongoDB session/transaction:")
code_block(doc, """const session = await mongoose.startSession();
try {
  session.startTransaction();
  const order = await createOrderFromCart({ ... , session });
  await session.commitTransaction();
  ...
} catch (error) {
  await session.abortTransaction();
  throw error;
} finally {
  session.endSession();
}""")
note(doc, "`createOrderFromCart` (helpers/createOrderFromCart.js) does, all within the session: getUserCart -> validateOrderItems -> buildOrderItemsSnapshot -> calculateOrderTotal -> Order.create -> updateProductStock -> clearCart. This is the single place an order is ever created, so stock is always decremented atomically with the order.")


# ===================== SECTION 4 =====================
b = reg("4. Data Models", 1)
heading(doc, "4. Data Models", 1, bookmark=b)
para(doc, "Nine Mongoose models live in `toy/backend/src/models/`. The two with the most business logic are `Order` (nested sub-schemas) and `User` (password hook). Relationships are expressed with `ObjectId` refs and populated on read.")

MODELS = [
    ("Product", [
        ("name", "String", "required, trimmed"),
        ("description", "String", "required"),
        ("price", "Number", "required, min 0"),
        ("discountPrice", "Number", "default 0"),
        ("category", "ObjectId -> Category", "required"),
        ("images", "[{url, public_id}]", "gallery"),
        ("stock", "Number", "default 0"),
        ("brand", "String", "default 'WonderFox'"),
        ("rating", "Number", "default 0 (0-5)"),
        ("numReviews", "Number", "default 0"),
        ("isFeatured", "Boolean", "default false"),
        ("isActive", "Boolean", "default true"),
        ("createdBy", "ObjectId -> User", "admin who created it"),
    ], "timestamps. `sellingPrice` (discount when >0 and <=price) is computed in the listing aggregation, NOT stored.",
     [("Category", "product.category -> Category"), ("User", "product.createdBy -> User"), ("Review", "1 product has many reviews"), ("Order", "order items reference product")]),
    ("User", [
        ("name", "String", "required, trimmed"),
        ("email", "String", "required, unique, lowercase"),
        ("password", "String", "required, min 6, `select: false`"),
        ("role", "String", "enum [customer, admin], default customer"),
        ("isActive", "Boolean", "default true"),
    ], "`pre('save')` hashes the password with bcryptjs (10 rounds) only when modified. `matchPassword(pw)` compares with bcrypt. Password is `select: false` so it never leaks into responses.",
     [("Order", "user -> orders"), ("Cart", "user -> cart (unique)"), ("Wishlist", "user -> wishlist (unique)")]),
    ("Order", [
        ("orderNumber", "String", "unique, indexed"),
        ("user", "ObjectId -> User", "required, indexed"),
        ("items", "[OrderItem]", "required, non-empty"),
        ("shippingAddress", "ShippingAddress", "required"),
        ("paymentMethod", "String", "enum [COD, RAZORPAY]"),
        ("paymentStatus", "String", "enum [PENDING, PAID, FAILED, REFUNDED], default PENDING"),
        ("paymentResult", "PaymentResult", "default {}"),
        ("orderStatus", "String", "enum [PENDING, CONFIRMED, PACKED, SHIPPED, DELIVERED, CANCELLED], default PENDING"),
        ("statusHistory", "[StatusHistory]", "default []"),
        ("itemsPrice", "Number", "required, min 0"),
        ("shippingPrice", "Number", "default 0"),
        ("taxPrice", "Number", "default 0"),
        ("totalPrice", "Number", "required, min 0"),
        ("paidAt", "Date", ""),
        ("deliveredAt", "Date", ""),
    ], "timestamps. Indexes: {user,createdAt:-1}, and a **sparse unique** index on `paymentResult.razorpayPaymentId` to prevent duplicate payments. Sub-schemas: `orderItemSchema` (product/name/image/qty/price/discountPrice; `_id:false`), `shippingAddressSchema` (Indian phone `^[6-9]\\d{9}$`, 6-digit postal `^\\d{6}$`, country fixed to 'India'; `_id:false`), `paymentResultSchema`, `statusHistorySchema` (status/updatedAt/updatedBy).",
     [("User", "order.user -> User"), ("Product", "order items reference product"), ("User", "statusHistory.updatedBy -> User")]),
    ("Category", [
        ("name", "String", "required, unique, trimmed"),
        ("description", "String", "trimmed"),
        ("image", "String", "default ''"),
        ("isActive", "Boolean", "default true"),
    ], "timestamps. A product belongs to one category; categories own an image (uploaded via multer).",
     [("Product", "category has many products")]),
    ("Cart", [
        ("user", "ObjectId -> User", "required, unique"),
        ("items", "[CartItem]", ""),
        ("totalAmount", "Number", "default 0, recomputed on save"),
    ], "timestamps. `cartItemSchema` = `{ product (ref Product), quantity (min 1, default 1), price }`. One cart per user.",
     [("User", "cart.user -> User (1:1)"), ("Product", "cart items reference product")]),
    ("Wishlist", [
        ("user", "ObjectId -> User", "required, unique, indexed"),
        ("items", "[{ product (ref Product) }]", ""),
    ], "timestamps. One wishlist per user (created lazily on first add/read). Items only store the product ref.",
     [("User", "wishlist.user -> User (1:1)"), ("Product", "wishlist items reference product")]),
    ("Blog", [
        ("title", "String", "required"),
        ("slug", "String", "required, unique, lowercase"),
        ("excerpt", "String", "required"),
        ("content", "String", "required"),
        ("coverImage", "String", "required on create"),
        ("author", "String", "default 'WonderFox Editorial'"),
        ("category", "String", "enum from BLOG_CATEGORIES"),
        ("tags", "[String]", "default []"),
        ("status", "String", "enum [draft, published], default draft"),
        ("publishedAt", "Date", "set when published, null when draft"),
    ], "timestamps. `category` is a free-text enum (Parenting Tips, Toy Guides, Learning, Gift Ideas, Kids Activities, News & Updates) - NOT a Category ref. Slug is made unique with a counter loop.",
     []),
    ("Review", [
        ("product", "ObjectId -> Product", "required"),
        ("user", "ObjectId -> Product/User", "required"),
        ("rating", "Number", "required, 1-5"),
        ("comment", "String", "maxlength 1000"),
    ], "timestamps. **Unique compound index** `{ product, user }` enforces one review per user per product (upsert on save).",
     [("Product", "review.product -> Product"), ("User", "review.user -> User")]),
    ("NewsletterSubscriber", [
        ("email", "String", "required, unique, lowercase, regex-validated"),
        ("isActive", "Boolean", "default true"),
        ("subscribedAt", "Date", "default now"),
        ("unsubscribedAt", "Date", "default null"),
        ("lastSentAt", "Date", "default null"),
    ], "timestamps. Soft-delete pattern: unsubscribing flips `isActive` and sets `unsubscribedAt` instead of deleting. Re-subscribing reactivates.",
     []),
]
for name, fields, notes, relations in MODELS:
    h3(doc, name + " model", bookmark="model_" + name)
    rows = [[f, t, n] for f, t, n in fields]
    table(doc, ["Field", "Type", "Notes"], rows)
    para(doc, "**Design notes:** " + notes)
    if relations:
        bul = ["`%s.%s` -> %s" % (name, f, t) if False else "%s -> %s" % (f, t) for f, t in relations]
        bullet(doc, ["**Relations:**"] + bul)


# ===================== SECTION 5 =====================
b = reg("5. API Reference", 1)
heading(doc, "5. API Reference", 1, bookmark=b)
para(doc, "All routes are mounted under `/api`. `protect` = JWT-authenticated; `admin` = JWT + role `admin`. Methods/paths come directly from the route files in `toy/backend/src/routes/`.")

API = [
    ("Authentication (`/api/auth`)",
     [("POST", "/register", "public", "register user -> 201 + token/user"),
      ("POST", "/login", "public", "login -> 200 + token/user"),
      ("GET", "/me", "protect", "current user profile")]),
    ("Products (`/api/products`)",
     [("GET", "/", "public", "getAllProducts - aggregation with search/category/price/sort/pagination"),
      ("GET", "/:id", "public", "getProductById (populated)"),
      ("POST", "/", "protect", "createProduct"),
      ("PUT", "/:id", "protect", "updateProduct"),
      ("DELETE", "/:id", "protect", "deleteProduct"),
      ("GET", "/:id/reviews", "public", "getProductReviews"),
      ("POST", "/:id/reviews", "protect", "saveProductReview (unique upsert)")]),
    ("Categories (`/api/categories`)",
     [("GET", "/", "public", "getAllCategories"),
      ("GET", "/:id", "public", "getCategoryById"),
      ("POST", "/", "protect + upload", "createCategory (image via multer)"),
      ("PUT", "/:id", "protect + upload", "updateCategory"),
      ("DELETE", "/:id", "protect", "deleteCategory")]),
    ("Cart (`/api/cart`) - all protected",
     [("GET", "/", "protect", "getCart"),
      ("POST", "/", "protect", "addToCart {productId, quantity}"),
      ("PATCH", "/:productId", "protect", "updateCartItem {quantity}"),
      ("DELETE", "/:productId", "protect", "removeCartItem"),
      ("DELETE", "/", "protect", "clearCart")]),
    ("Wishlist (`/api/wishlist`) - all protected",
     [("POST", "/", "protect", "addToWishlist {productId}"),
      ("GET", "/", "protect", "getWishlist"),
      ("DELETE", "/", "protect", "clearWishlist"),
      ("DELETE", "/:productId", "protect", "removeFromWishlist")]),
    ("Orders (`/api/orders`) - customer protected",
     [("POST", "/", "protect", "placeCodOrder"),
      ("POST", "/razorpay", "protect", "createRazorpayOrder"),
      ("POST", "/verify-payment", "protect", "verifyPayment -> order"),
      ("GET", "/", "protect", "getMyOrders (latest first)"),
      ("GET", "/:id", "protect", "getOrderById (owner check)"),
      ("PATCH", "/:id/cancel", "protect", "cancelOrder (PENDING/CONFIRMED only)")]),
    ("Orders (`/api/orders`) - admin only",
     [("GET", "/admin", "admin", "getAllOrders"),
      ("GET", "/admin/:id", "admin", "getOrderByIdForAdmin"),
      ("PATCH", "/admin/:id/status", "admin", "updateOrderStatus (state-machine guard)")]),
    ("Upload (`/api/upload`)",
     [("POST", "/", "protect + upload", "uploadImage -> {url, public_id}")]),
    ("Blogs (`/api/blogs`)",
     [("GET", "/", "public", "getPublicBlogs (published)"),
      ("GET", "/:slug", "public", "getBlogBySlug"),
      ("GET", "/admin", "admin", "getAdminBlogs (search/status/category)"),
      ("POST", "/", "admin", "createBlog"),
      ("PUT", "/:id", "admin", "updateBlog"),
      ("DELETE", "/:id", "admin", "deleteBlog")]),
    ("Newsletter (`/api/newsletter`)",
     [("POST", "/subscribe", "public", "subscribeToNewsletter"),
      ("GET", "/subscribers", "admin", "getSubscribers"),
      ("DELETE", "/subscribers/:id", "admin", "deactivateSubscriber"),
      ("POST", "/send", "admin", "sendNewsletterToSubscribers")]),
    ("Contact (`/api/contact`)",
     [("POST", "/", "public", "submitContactMessage -> email via nodemailer")]),
    ("Dashboard (`/api/dashboard`)",
     [("GET", "/", "protect", "getDashboardStats - counts + revenue + low stock")]),
    ("Users (`/api/users`) - admin only",
     [("GET", "/admin", "admin", "getAdminUsers (with $lookup order summary)"),
      ("GET", "/admin/:id", "admin", "getAdminUserById (orders + spend)")]),
]
for title, endpoints in API:
    subheading(doc, title, bookmark="api_" + title.replace("/", "").split()[0].lower())
    rows = [[m, p, a, d] for m, p, a, d in endpoints]
    table(doc, ["Method", "Path", "Auth", "Description"], rows,
          col_widths=[Cm(1.6), Cm(4.2), Cm(2.0), Cm(7.2)])


# ===================== SECTION 6 =====================
b = reg("6. Feature Deep-Dives", 1)
heading(doc, "6. Feature Deep-Dives", 1, bookmark=b)

subheading(doc, "6.1 Authentication flow (JWT in localStorage)", bookmark="6_1")
para(doc, "Register and login both return `{ token, user }`. The token is `jwt.sign({ id }, JWT_SECRET, { expiresIn: JWT_EXPIRE })`. The frontend keeps it in `localStorage` (NOT an HttpOnly cookie). The shared Axios `api` instance injects `Authorization: Bearer <token>` on every request via a request interceptor, and on a 401 response clears `token`/`user` from storage and hard-redirects to `/login?redirect=<original-path>` so the shopper does not lose their cart intent:")
code_block(doc, """// frontend/src/api/axios.js
api.interceptors.response.use(
  (response) => response,
  (error) => {
    if (error.response?.status === 401) {
      localStorage.removeItem("token");
      localStorage.removeItem("user");
      window.location.href = `/login?redirect=${encodeURIComponent(window.location.pathname)}`;
    }
    return Promise.reject(error);
  },
);""")
note(doc, "Security trade-off: bearer-in-localStorage is convenient and stateless but is vulnerable to XSS. An HttpOnly cookie would mitigate XSS but is NOT used here (see Limitations).")

subheading(doc, "6.2 Guest cart merge strategy", bookmark="6_2")
para(doc, "Unauthenticated shoppers build a cart in `localStorage` under key `guestCart` (CartContext writes/normalizes it via `normalizeGuestItem`). `CartContext.addItem` branches on the presence of a token: no token -> write to `localStorage`; token present -> call `POST /cart`. After login, `fetchCart()` swaps in the server-side cart. This means a guest's items are preserved client-side across the login boundary, but are **not** automatically merged into the user's server cart - the cart resets to the logged-in user's persisted cart once a token exists.")

subheading(doc, "6.3 Cart total & item validation", bookmark="6_3")
para(doc, "`calculateCartTotal(items)` simply sums `item.price * item.quantity`. Quantities are validated with `validateQuantity` (positive integer). Each add/update checks `quantity > product.stock` and throws `400 Insufficient stock`. Stock is only actually decremented at order creation time, so the cart stays editable until checkout.")
code_block(doc, """const calculateCartTotal = (items) =>
  items.reduce((total, item) => total + item.price * item.quantity, 0);""")

subheading(doc, "6.4 Checkout & the two payment flows", bookmark="6_4")
para(doc, "Both flows funnel through the **same** `createOrderFromCart` helper, so orders are created identically regardless of payment method - only the payment status differs:")
numbered(doc, [
    "**Cash on Delivery** - `placeCodOrder` creates the order with `paymentMethod: 'COD'`, `paymentStatus: 'PENDING'`, `orderStatus: 'PENDING'` and clears the cart.",
    "**Razorpay** - two-stage: (1) `createRazorpayOrder` builds a Razorpay Order from the cart total plus 5% GST, amount in paise; (2) the frontend confirms with Razorpay, then calls `verifyPayment` which (a) recomputes the HMAC signature, (b) rejects duplicate `razorpayPaymentId`s via the unique sparse index, and (c) inside a transaction calls `createOrderFromCart` with `paymentMethod: 'RAZORPAY'`, `paymentStatus: 'PAID'`, `orderStatus: 'CONFIRMED'`.",
])
code_block(doc, """// payment.service.js - signature verification
const body = `${razorpayOrderId}|${razorpayPaymentId}`;
const expected = crypto.createHmac("sha256", process.env.RAZORPAY_KEY_SECRET).update(body).digest("hex");
return expected === razorpaySignature;""")
note(doc, "Pricing detail: `calculateOrderTotal` sets `itemsPrice = sum(price*qty)`, `shippingPrice = itemsPrice >= 150 ? 0 : 100`, `taxPrice = 0`, `totalPrice = itemsPrice + shippingPrice + tax`. The Razorpay path additionally adds 5% GST to the *amount sent to Razorpay*, but the stored Order `totalPrice` uses `calculateOrderTotal` (GST excluded) - a small inconsistency worth flagging in review.")

subheading(doc, "6.5 Order number generation", bookmark="6_5")
para(doc, "`generateOrderNumber` produces `WF-<YYYYMMDD>-<6 hex chars>` using `crypto.randomBytes(3)`. It is NOT guaranteed unique by itself (no loop), but collisions are implausible and `orderNumber` has a unique index - acceptable for this scale.")
code_block(doc, """const date = new Date().toISOString().slice(0, 10).replace(/-/g, "");
const random = crypto.randomBytes(3).toString("hex").toUpperCase();
return `WF-${date}-${random}`;""")

subheading(doc, "6.6 Order status state machine", bookmark="6_6")
para(doc, "`order.service.js` enforces allowed transitions server-side, so the admin UI cannot jump to an illegal status:")
bullet(doc, [
    "PENDING -> CONFIRMED | CANCELLED",
    "CONFIRMED -> PACKED | CANCELLED",
    "PACKED -> SHIPPED",
    "SHIPPED -> DELIVERED",
    "DELIVERED -> (terminal)",
    "CANCELLED -> (terminal)",
])
para(doc, "A customer may only cancel from PENDING or CONFIRMED. Every transition is appended to `statusHistory` with the acting `updatedBy` user, and delivering sets `deliveredAt` automatically.")

subheading(doc, "6.7 Stock management", bookmark="6_7")
para(doc, "Stock is checked in two places and decremented in one: (1) `validateOrderItems` verifies `product.stock >= qty` before an order is built; (2) `updateProductStock` runs a `Product.bulkWrite` of `$inc: { stock: -quantity }` operations, all within the same session transaction that created the order and cleared the cart. This prevents overselling under concurrent checkouts.")

subheading(doc, "6.8 Reviews", bookmark="6_8")
para(doc, "A customer can write **one** review per product: the unique compound index `{ product, user }` and `findOneAndUpdate({...},{ upsert: true })` make `/reviews` an upsert. Saving a review recalculates the product's `rating` (avg) and `numReviews` via a `$group` aggregation, so the stored `rating` is always derived from real reviews.")

subheading(doc, "6.9 Blog CMS", bookmark="6_9")
para(doc, "Blogs use a free-text `category` enum (not a Category ref) and are addressed by `slug` for public URLs. `slugify` strips non-alphanumerics to kebabs, and `buildUniqueSlug` loops with a counter until the slug is free. Drafts stay hidden (`status: 'draft'`; `publishedAt` null) and only `published` posts are served publicly.")

subheading(doc, "6.10 Newsletter & contact email", bookmark="6_10")
para(doc, "Both reuse a shared `nodemailer` SMTP transport. `subscribeToNewsletter` normalizes + regex-validates the email and **reactivates** a previously unsubscribed address (409 only if already active). `sendNewsletterToSubscribers` HTML-emails active subscribers one-by-one, records `lastSentAt` per send, and counts successes/failures. The contact form sanitizes input with `escapeHtml` and `replyTo`s the shopper's address.")

subheading(doc, "6.11 Image uploads", bookmark="6_11")
para(doc, "Multer (`uploadMiddleware.js`) buffers uploads to **memory** (5 MB cap). The upload service base64-encodes the buffer and streams it to Cloudinary's `wonderfox` folder, returning `{ url, public_id }`. The same middleware is reused by `POST /upload` and the category create/update routes.")

subheading(doc, "6.12 Dashboard statistics", bookmark="6_12")
para(doc, "`getDashboardStats` runs lightweight queries: `Product.countDocuments()`, recent products, low-stock products (stock <= 5), category count, order count, and `totalRevenue` (sum of all `order.totalPrice`, including CANCELLED orders - a possible refinement). No aggregation fancy math, deliberately cheap.")

subheading(doc, "6.13 Limitations & planned enhancements", bookmark="6_13")
bullet(doc, [
    "JWT in localStorage is XSS-prone; consider rotating refresh tokens in an HttpOnly cookie.",
    "express-validator is installed but unused - introduce schema validation on routes.",
    "`getProfile` client service calls `/auth/profile` but the route is `/auth/me` - minor endpoint drift.",
    "Frontend `AuthContext.jsx` and `WishlistContext.jsx` exist as empty files - auth & wishlist state are handled inline/locally.",
    "Total revenue counts cancelled orders; filter those out.",
    "`react-refresh`/`@vitejs/plugin-react` are declared as devDependencies of the frontend but are not actually in `frontend/package.json` (lint config references them).",
])
planned(doc, "Add Stripe as an alternative payment provider; add server-side rendering or pre-rendered PDPs for SEO.")


# ===================== SECTION 7 =====================
b = reg("7. Frontend Architecture", 1)
heading(doc, "7. Frontend Architecture", 1, bookmark=b)
para(doc, "Two independent React 19 + Vite 8 SPAs share an architecture but not a codebase: the **customer storefront** (`toy/frontend`) and the **admin panel** (`toy/admin`). Both consume the same backend API.")

subheading(doc, "7.1 Project scaffold", bookmark="7_1")
para(doc, "`main.jsx` renders `AppProviders` (which wraps the app in `CartProvider` and a global `<Toaster>`) inside `BrowserRouter`. `App.jsx` defines a `MainLayout` (header + footer via `Outlet`) for public pages; `/login` and `/register` are rendered without the layout; `*` falls back to `NotFound`.")
bullet(doc, [
    "Pages: Home, Collection, ProductDetails, Cart, Checkout, Wishlist, MyOrders, OrderSuccess, Blog, BlogDetail, About, Contact, Login, Register, Profile, NotFound.",
    "Layouts: MainLayout (public), AdminLayout (admin SPA).",
    "Reusable UI: Button (primary/secondary), Input, Card, Badge, SectionTitle, Container, SectionHeading, RatingStars, FeatureCard, CategoryCard, ProductCard, TestimonialCard, ProfileDropdown.",
])

subheading(doc, "7.2 State management", bookmark="7_2")
para(doc, "State is intentionally minimal - no Redux/Zustand:")
bullet(doc, [
    "**Cart** - `CartContext` (the only React Context actually used). Holds cart items/total, exposes `addItem/updateItem/removeItem/clearCart`, and bridges guest (`localStorage` `guestCart`) and authenticated carts.",
    "**Auth & Wishlist** - `AuthContext.jsx` and `WishlistContext.jsx` exist but are **empty files**. Auth is resolved at request time (axios interceptor reads `localStorage.token`), and the Wishlist page talks to the API directly.",
    "**Local UI** - `useState`/`useEffect` per component.",
    "**Routing state** - query strings and `useSearchParams`/`useParams`/`useNavigate` read URL state.",
])
note(doc, "The previous search log shows `useContext`, `createContext`, `useParams`, `useNavigate` and `useSearchParams` appearing in the codebase - `createContext` is used only by CartContext; the other React-Router hooks are used in route-aware pages (ProductDetails uses `useParams` for `:id`).")

subheading(doc, "7.3 API layer", bookmark="7_3")
para(doc, "Every service imports the shared `api` Axios instance (`src/api/axios.js`), which sets `baseURL` from `VITE_API_BASE_URL` (fallback to the Render deploy URL), JSON headers, a request interceptor for the Bearer token, and a 401 response interceptor that logs the user out. Services are thin wrappers over `api`:")
bullet(doc, [
    "auth.service - loginUser/registerUser/getProfile",
    "product.service - getProducts/getProduct/getProductReviews/saveProductReview",
    "cart.service - getCart/addToCart/updateCart/removeFromCart/clearCart",
    "order.service - createCODOrder/createRazorpayOrder/verifyPayment/getMyOrders",
    "blog, category, wishlist, etc.",
])
code_block(doc, """export const createRazorpayOrder = async () => {
  const { data } = await api.post("/orders/razorpay");
  return data.data;
};""")

subheading(doc, "7.4 Styling & design system", bookmark="7_4")
para(doc, "Tailwind CSS 4 is imported via `@import \"tailwindcss\"` in `index.css`, and WonderFox's brand palette is declared as CSS custom properties inside an `@theme` block: `--color-primary` (#3B82F6), `--color-secondary`/`--color-accent` (#FF7A59/#10B981), plus surface/text/status tokens and `--radius-card` (24px) / `--radius-button` (16px). Utility classes are used throughout (e.g. the Footer uses Tailwind's `orange-*` palette and `grid` + `sm:/md:` responsive breakpoints).")

subheading(doc, "7.5 Interaction & animation", bookmark="7_5")
para(doc, "Framer Motion powers home-section reveals (`motion.div` + `AnimatePresence`) and the mobile filter drawer slide. `react-icons/fa6` provides the Footer social icons; `lucide-react` supplies in-app icons. `react-hot-toast` gives non-blocking feedback (add-to-cart, checkout result, errors).")


# ===================== SECTION 8 =====================
b = reg("8. Interview Q&A", 1)
heading(doc, "8. Interview Q&A", 1, bookmark=b)

QA = [
    ("What architectural pattern does the backend follow?",
     "Route -> Controller -> Service -> Model. Controllers are thin (asyncHandler + delegate); services own business logic and return a uniform ApiResponse/ApiError; Mongoose models own persistence. This keeps business rules testable and independent of Express."),
    ("Why is the JWT kept in localStorage instead of an HttpOnly cookie?",
     "To stay stateless and avoid CSRF (the API is CORS + Bearer-only, and `cookie-parser` is even unused). Trade-off: localStorage is readable from JS, so XSS could steal the token. The mitigation in this code is the axios 401 handler that clears storage, but a hardened version would use an HttpOnly refresh-token cookie."),
    ("How are passwords protected?",
     "Stored as bcryptjs hashes via a `User.pre('save')` hook (10 rounds), and the field is `select: false` so Mongoose omits it from queries by default. `matchPassword` compares with `bcrypt.compare`."),
    ("How does the cart work for guests vs logged-in users?",
     "CartContext branches on `localStorage.token`. Guests write to `localStorage.guestCart` (normalized via `normalizeGuestItem`). Authenticated users call the `/cart` API which persists to MongoDB. `fetchCart` loads the server cart on login."),
    ("How is overselling prevented?",
     "At two layers: `validateOrderItems` checks `stock >= qty` before building an order, and `updateProductStock` runs `Product.bulkWrite` with `$inc: { stock: -qty }` inside the same session transaction that creates the order and clears the cart - so the decrement is atomic with the order write."),
    ("How is a double-charge / duplicate order prevented during payment?",
     "Two guards: (1) an explicit `Order.findOne({ 'paymentResult.razorpayPaymentId': ... })` check, and (2) a **sparse unique index** on that field, which turns any race into a 11000 duplicate-key error caught and converted to 409. Razorpay signature is HMAC-verified server-side before any order is written."),
    ("Why is the product listing an aggregation pipeline rather than `find` + sort?",
     "Because `sellingPrice` (discount vs list) is computed with `$cond`/`$addFields` and then filtered by `minPrice`/`maxPrice` and sorted by that computed value - something Mongoose `find` sorting cannot do in one pass. `$facet` returns both the page of data and the metadata (total count/pagination) in a single round-trip."),
    ("How is the order status protected from invalid transitions?",
     "`updateOrderStatus` uses an explicit `allowedTransitions` map. Anything not permitted returns 400. This is enforced server-side regardless of what the admin frontend sends."),
    ("What does `express-validator` do in this project?",
     "Nothing - it is declared in package.json but never imported. All validation is manual (validateQuantity, validateObjectId, getValidProduct, validateOrderItems, validateBlogPayload)."),
    ("Why does the frontend use `cookie-parser` on the backend if auth is token-based?",
     "It doesn't - `cookie-parser` is `app.use(cookieParser())` but no route reads `req.cookies` for auth. The dependency is vestigial; auth is Bearer-token only. (This is explicitly flagged as a limitation.)"),
    ("How is the blog SEO-friendly?",
     "Public posts are fetched by `slug` (`/api/blogs/:slug`), drafts are hidden (`status:'draft'`, `publishedAt:null` when draft), and the frontend ProductDetails page renders SEO metadata + JSON-LD structured data (noted in the feature list)."),
    ("How would you scale this beyond a single node?",
     "Run multiple backend instances behind a load balancer, switch the stateless JWT to a distributed cache for revocation lists, add a CDN in front of Cloudinary + the static SPA, and index MongoDB on the hot queries (orders by user+date, products by category, low stock). The aggregation pipeline is already the right shape for pagination."),
]
for question, answer in QA:
    para(doc, "Q. " + question)
    _bold_label(doc, "A.", answer)
    para(doc, "")

def _bold_label(doc, label, text):
    par = doc.add_paragraph()
    run = par.add_run(label)
    run.bold = True
    _add_rich_runs(par, " " + text)
    return par

# ----------------------------------------------------------------------------
# Finalize: build the clickable TOC, then save.
# ----------------------------------------------------------------------------

build_toc(doc, TOC_MARKER)
footer_with_page_numbers(doc)  # re-affirm on the (now stable) section

doc.save(OUTPUT)
print("Wrote", OUTPUT)
